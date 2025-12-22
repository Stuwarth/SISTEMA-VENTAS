import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = tcPr.makeelement(qn('w:tcMar'))
        tcPr.append(tcMar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{side}'))
        if node is None:
            node = tcMar.makeelement(qn(f'w:{side}'))
            tcMar.append(node)
        node.set(qn('w:w'), str(int(value * 567))) 
        node.set(qn('w:type'), 'dxa')

def convert():
    md_file = 'CAPITULO_I_final.md'
    if not os.path.exists(md_file):
        print(f"Error: {md_file} no existe.")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # 1. Page Setup (Letter size and Margins)
    section = doc.sections[0]
    section.page_height = Cm(27.94) 
    section.page_width = Cm(21.59)  
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)

    # 2. Main Styles (Arial, 12pt, 1.5 spacing, Justified)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        # Chapter Title: **CAPÍTULO X**
        if re.match(r'^\*\*CAPÍTULO [IVXLCDM]+\*\*$', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace('**', '').upper())
            run.font.name = 'Arial Black'
            run.font.bold = True
            run.font.size = Pt(36)
            
            i += 1
            if i < len(lines) and lines[i].strip():
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Remove markdown bold syntax and convert to upper
                title_text = lines[i].strip().replace('**', '').upper()
                run2 = p2.add_run(title_text)
                run2.font.name = 'Arial Black'
                run2.font.bold = True
                run2.font.size = Pt(36)
            i += 1
            continue

        # Subtitle: **X.X ...** or **X ...**
        elif re.match(r'^\*\*\d+(\.\d+)*\s.*?\*\*$', line) or re.match(r'^\d+\.\s.*$', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('**', '')
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(12)
            i += 1
            continue

        # Tables: | ... |
        elif line.startswith('|'):
            # Check for table title above if exists in previous paragraph or line starts with TABLA
            # Since we process line by line, if the current is a table, look for previous if it was a title
            
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 2:
                # Potential table logic
                rows_data = []
                for idx, tline in enumerate(table_lines):
                    if '---' in tline and idx == 1: continue
                    cols = [c.strip().replace('**', '') for c in tline.split('|')[1:-1]]
                    rows_data.append(cols)
                
                if rows_data:
                    table = doc.add_table(rows=len(rows_data), cols=max(len(r) for r in rows_data))
                    table.style = 'Table Grid'
                    for r_idx, r_data in enumerate(rows_data):
                        row = table.rows[r_idx]
                        for c_idx, c_text in enumerate(r_data):
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.text = c_text
                                cell_p = cell.paragraphs[0]
                                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                                run = cell_p.runs[0] if cell_p.runs else cell_p.add_run(c_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                                if r_idx == 0:
                                    run.font.bold = True
            
            # Check for "Fuente:" below
            if i < len(lines) and "Fuente:" in lines[i]:
                p_f = doc.add_paragraph()
                p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_f = p_f.add_run(lines[i].strip().replace('*', ''))
                run_f.font.name = 'Arial'
                run_f.font.italic = True
                run_f.font.size = Pt(10)
                i += 1
            continue

        # Diagram/Placeholder
        elif "![PONER DIAGRAMA AQUI" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.italic = True
            i += 1
            continue

        # Normal text or anything else
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Match markdown bold in line for internal runs
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.font.bold = True
                else:
                    p.add_run(part)
            i += 1

    output_name = 'PROYECTO_FINAL_SIS_VENTAS_STANDAR.docx'
    doc.save(output_name)
    print(f"Archivo generado: {output_name}")

if __name__ == "__main__":
    convert()
