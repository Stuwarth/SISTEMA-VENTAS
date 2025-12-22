import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(doc):
    # This adds a page number to the footer, bottom right
    # And a thin line above it
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add thin line above page number (border top of paragraph)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4') # Approx 0.5pt
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), 'auto')
        pBdr.append(top)
        pPr.append(pBdr)

        # Page number field
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

def convert():
    md_file = 'CAPITULO_I_final.md'
    if not os.path.exists(md_file):
        print(f"Error: {md_file} no existe.")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # 1. Page Design: Carta (Letter), Margins: L 3.5, R 2.5, T 2.5, B 2.5
    section = doc.sections[0]
    section.page_height = Cm(27.94) 
    section.page_width = Cm(21.59)  
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)

    # 2. Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_page_number(doc)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        # Chapter Title (Level 1)
        # Matches **CAPÍTULO X**
        if re.match(r'^\*\*CAPÍTULO [IVXLCDM]+\*\*$', line):
            # Line 1: CAPÍTULO X
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace('**', '').upper())
            run.font.name = 'Arial Black'
            run.font.bold = True
            run.font.size = Pt(36)
            
            # Line 2: TITLE
            i += 1
            while i < len(lines) and not lines[i].strip(): i += 1 # skip empty
            if i < len(lines):
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_text = lines[i].strip().replace('**', '').upper()
                run2 = p2.add_run(title_text)
                run2.font.name = 'Arial Black'
                run2.font.bold = True
                run2.font.size = Pt(36)
            i += 1
            continue

        # Subtitles / Headers (1.1, 1.2, etc.)
        elif re.match(r'^\*\*\d+(\.\d+)*\s.*?\*\*$', line) or re.match(r'^\d+\.\d+\s.*$', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('**', '')
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(12)
            i += 1
            continue

        # Table Titles: **TABLA X: NOMBRE** (Should be above table)
        elif line.startswith('**TABLA') or line.startswith('**ILUSTRACIÓN'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace('**', '').upper())
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(12)
            i += 1
            continue

        # Tables: | ... |
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 2:
                rows_data = []
                for idx, tline in enumerate(table_lines):
                    if '---' in tline and idx == 1: continue
                    cols = [c.strip().replace('**', '') for c in tline.split('|')[1:-1]]
                    rows_data.append(cols)
                
                if rows_data:
                    table = doc.add_table(rows=len(rows_data), cols=max(len(r) for r in rows_data))
                    table.style = 'Table Grid'
                    for r_idx, r_data in enumerate(rows_data):
                        row = table.rows[r_idx]
                        for c_idx, c_text in enumerate(r_data):
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.text = c_text
                                cell_p = cell.paragraphs[0]
                                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                                run = cell_p.runs[0] if cell_p.runs else cell_p.add_run(c_text)
                                run.font.name = 'Arial'
                                run.font.size = Pt(10) # 10-11pt as requested
                                if r_idx == 0:
                                    run.font.bold = True
            continue

        # Source / Fuente below tables/images
        elif "Fuente:" in line or "*Fuente:*" in line or "*FUENTE:*" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('*', '').replace('_', '')
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Diagram Placeholders
        elif "![PONER DIAGRAMA AQUI" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name = 'Arial'
            run.font.italic = True
            run.font.size = Pt(12)
            i += 1
            continue

        # Lists
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(line[2:].replace('**', ''))
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            # Indent 1.27cm as per specs
            p.paragraph_format.left_indent = Cm(1.27)
            i += 1
            continue

        # Regular Paragraphs
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Process internal bold
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.font.bold = True
                else:
                    p.add_run(part)
            i += 1

    output_name = 'DOC_SISTEMA_VENTAS_BOLIVIA_2025.docx'
    doc.save(output_name)
    print(f"Éxito: Documento '{output_name}' generado con formato estricto.")

if __name__ == "__main__":
    convert()
